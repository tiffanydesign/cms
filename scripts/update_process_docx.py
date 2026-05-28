"""
Update Process.docx with corrections and new section for script 13.

Changes:
1. Fix §5.1 TTC contrasts vs 0 Hz: all are non-significant (p = .426), NOT significant
2. Update CSV count: 59 → 62 (in §1.2, §7.1, footer)
3. Update PNG count: 17 → 21 (in §1.2, §7.2, footer)
4. Add new section 3.15 for 13_post_survey_descriptives.py
5. Update 10_plots.R description to include 2 additional QQ figures
6. Update script count in footer: 14 R + 1 Python = add Python note
"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import copy
import sys, io

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

DOCX_PATH = r"C:\Users\50560\Desktop\cms_analysis\writing_reference\Process.docx"

doc = Document(DOCX_PATH)

# ─── Helper: replace text in a paragraph while preserving run formatting ───────
def replace_para_text(para, new_text):
    """Replace entire paragraph content with new_text, keeping first run's formatting."""
    if not para.runs:
        para.add_run(new_text)
        return
    # Keep first run's font/bold/size, clear rest
    first_run = para.runs[0]
    first_run.text = new_text
    for run in para.runs[1:]:
        run.text = ""


def para_contains(para, text):
    return text in para.text


# ─── Print paragraph index for debugging (run once, then comment out) ─────────
# for i, p in enumerate(doc.paragraphs):
#     print(f"[{i:03d}] {repr(p.text[:80])}")

paragraphs = doc.paragraphs

# ─── 1. Fix §5.1 TTC planned contrasts (P208, P209, P210) ────────────────────
# The three bullet items under 5.1 currently say "显著" (significant)
# but data shows all are p=.426 (non-significant after Holm correction)

fix_map = {
    "8.33 Hz vs 0 Hz：显著（p < .05，Holm 校正）":
        "8.33 Hz vs 0 Hz：不显著（p = .426，Holm 校正；t(24) = −1.311，95% CI [−0.644, 0.209] s）",
    "12.5 Hz vs 0 Hz：显著（p < .01）":
        "12.5 Hz vs 0 Hz：不显著（p = .426，Holm 校正；t(24) = −1.518，95% CI [−0.768, 0.198] s）",
    "25 Hz vs 0 Hz：显著（p < .001）":
        "25 Hz vs 0 Hz：不显著（p = .426，Holm 校正；t(24) = +1.225，95% CI [−0.236, 0.664] s）",
}

# Also need to find the 5.1 heading paragraph to insert a correction note after the bullets
section51_idx = None
bullets_51 = []

for i, p in enumerate(paragraphs):
    for old_text, new_text in fix_map.items():
        if old_text in p.text:
            replace_para_text(p, p.text.replace(old_text, new_text))
            print(f"  Fixed §5.1 bullet at P{i:03d}")
            bullets_51.append(i)
    if "5.1  TTC — 频率主效应" in p.text:
        section51_idx = i
        print(f"  Found §5.1 heading at P{i:03d}")

# Insert correction note after the last §5.1 bullet
# We'll add a paragraph after the last bullet with a bold note
if bullets_51:
    last_bullet_idx = max(bullets_51)
    last_bullet_para = paragraphs[last_bullet_idx]
    # Insert a new paragraph after last bullet by manipulating XML
    new_para = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'Normal')
    pPr.append(pStyle)
    new_para.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    b = OxmlElement('w:b')
    rPr.append(b)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = ("注：Step 1 整体频率主效应显著，F(1.83, 43.91) = 5.094，p = .012，η²p = .175（中等）。"
              "但三个计划对比（各闪烁频率 vs 0 Hz 稳定条件，Holm 校正）均未达显著水平（所有 p = .426），"
              "表明频率效应由闪烁条件之间的差异驱动——尤其是 25 Hz TTC 显著长于 8.33 Hz 和 12.5 Hz——"
              "而非单个闪烁频率与稳定基线的差异。")
    r.append(t)
    new_para.append(r)
    # Insert after last bullet paragraph
    last_bullet_para._element.addnext(new_para)
    print(f"  Inserted §5.1 correction note after P{last_bullet_idx:03d}")

# ─── 2. Fix CSV count 59 → 62 ──────────────────────────────────────────────────
csv_replacements = [
    ("共 59 个）", "共 62 个）"),
    ("共 59 个", "共 62 个"),
    ("CSV 输出：59 个", "CSV 输出：62 个"),
]

for i, p in enumerate(paragraphs):
    for old, new in csv_replacements:
        if old in p.text:
            replace_para_text(p, p.text.replace(old, new))
            print(f"  Fixed CSV count at P{i:03d}: '{old}' → '{new}'")

# ─── 3. Fix PNG count 17 → 21 ──────────────────────────────────────────────────
png_replacements = [
    ("共 17 个 PNG）", "共 21 个 PNG）"),
    ("共 17 个 PNG", "共 21 个 PNG"),
    ("图表输出：17 个", "图表输出：21 个"),
]

for i, p in enumerate(paragraphs):
    for old, new in png_replacements:
        if old in p.text:
            replace_para_text(p, p.text.replace(old, new))
            print(f"  Fixed PNG count at P{i:03d}: '{old}' → '{new}'")

# ─── 4. Update 10_plots.R to mention 2 additional QQ figures ──────────────────
for i, p in enumerate(paragraphs):
    if "3.11  10_plots.R — 所有主图生成" in p.text:
        plots_heading_idx = i
        print(f"  Found 10_plots.R heading at P{i:03d}")
    if "生成论文中使用的 14 个主要图表" in p.text:
        replace_para_text(p, p.text.replace(
            "生成论文中使用的 14 个主要图表",
            "生成论文中使用的 16 个主要图表"
        ))
        print(f"  Fixed 10_plots.R figure count at P{i:03d}")

# ─── 5. Update 총 script count in footer ──────────────────────────────────────
for i, p in enumerate(paragraphs):
    if "总脚本数：14" in p.text:
        replace_para_text(p, p.text.replace(
            "总脚本数：14",
            "总脚本数：15（14 个 R 脚本 + 1 个 Python 脚本）"
        ))
        print(f"  Fixed script count at P{i:03d}")
    if "生成时间：2026-05-27" in p.text:
        replace_para_text(p, p.text.replace(
            "生成时间：2026-05-27",
            "生成时间：2026-05-28（更新）"
        ))
        print(f"  Fixed date at P{i:03d}")

# ─── 6. Add §3.15 for script 13 (post-survey analysis) ───────────────────────
# Find section 3.14 heading to insert after it (and its content)
# Strategy: find the §3.14 section's last paragraph, then insert new section after

section314_heading_idx = None
for i, p in enumerate(paragraphs):
    if "3.14  12_ancova_covariates.R" in p.text:
        section314_heading_idx = i
        print(f"  Found §3.14 heading at P{i:03d}")
        break

# Find the next Heading 1 or Heading 2 after §3.14 to determine insertion point
if section314_heading_idx is not None:
    insert_before_idx = None
    for i in range(section314_heading_idx + 1, len(paragraphs)):
        if paragraphs[i].style.name in ('Heading 1', 'Heading 2'):
            # Check if this is chapter 4 (next top-level section)
            if "第四章" in paragraphs[i].text or paragraphs[i].style.name == 'Heading 1':
                insert_before_idx = i
                print(f"  Will insert §3.15 before P{i:03d}: {paragraphs[i].text[:60]}")
                break

    if insert_before_idx is not None:
        # Reference paragraph to insert before
        ref_para = paragraphs[insert_before_idx]
        ref_elem = ref_para._element

        def make_para(style_name, text, bold=False):
            """Create a paragraph XML element."""
            p_elem = OxmlElement('w:p')
            pPr = OxmlElement('w:pPr')
            pStyle = OxmlElement('w:pStyle')
            pStyle.set(qn('w:val'), style_name)
            pPr.append(pStyle)
            p_elem.append(pPr)
            r = OxmlElement('w:r')
            if bold:
                rPr = OxmlElement('w:rPr')
                b = OxmlElement('w:b')
                rPr.append(b)
                r.append(rPr)
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = text
            r.append(t)
            p_elem.append(r)
            return p_elem

        def make_bullet(text):
            """Create a List Bullet paragraph."""
            p_elem = OxmlElement('w:p')
            pPr = OxmlElement('w:pPr')
            pStyle = OxmlElement('w:pStyle')
            pStyle.set(qn('w:val'), 'ListBullet')
            pPr.append(pStyle)
            p_elem.append(pPr)
            r = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = text
            r.append(t)
            p_elem.append(r)
            return p_elem

        # Build paragraphs for §3.15 in READING ORDER (heading first → content last).
        # addprevious(ref) places each element right before ref, so iterating in
        # reading order produces the correct document sequence.
        new_sections = [
            # Heading FIRST
            make_para('Heading2',
                '3.15  13_post_survey_descriptives.py — 实验后问卷描述性分析（Python）'),
            # Description
            make_para('Normal',
                '本脚本对实验结束后填写的问卷进行描述性统计分析，涵盖驾驶困难度、心理负荷、努力程度、'
                'CMS 运动清晰度、闪烁干扰、视觉不适、决策自信度、闪烁可察觉度和闪烁对决策影响 9 个维度。'
                '分析结果可作为操纵检验（确认闪烁被感知）和主观评价量表的聚合效度证据。'),
            # Input data
            make_para('Normal', '输入数据：'),
            make_bullet('1. Experiment_Result/Google Forms/Post-Experiment Questionnaire  (Responses).xlsx'),
            make_bullet('注：文件名含双空格，脚本使用精确字符串匹配读取'),
            # Ordinal scales
            make_para('Normal', '9 题有序量表（括号内为选项数量）：'),
            make_bullet('Q1 驾驶道路变换困难度（6 级：Easy → Very difficult）'),
            make_bullet('Q2 心理负荷（6 级：Low → Very high）'),
            make_bullet('Q3 付出努力（6 级：Little → Very much）'),
            make_bullet('Q4 CMS 运动清晰度（6 级：Very unclear → Clear）'),
            make_bullet('Q5 闪烁干扰感（5 级：Slightly → Extremely）'),
            make_bullet('Q6 视觉不适 / 眼疲劳（5 级：Hardly any → Severe）'),
            make_bullet('Q7 决策自信度（5 级：Not confident at all → Confident）'),
            make_bullet('Q8 闪烁差异可察觉度（4 级：Slightly noticeable → Extremely noticeable）'),
            make_bullet('Q9 闪烁对决策的影响（5 级：Hardly at all → Significantly）'),
            # Method
            make_para('Normal', '统计方法：'),
            make_bullet('描述性统计：M、SD、中位数、IQR（Pandas / SciPy）'),
            make_bullet('众数：每题出现频率最高的回答选项'),
            make_bullet('频率分布：每个回答选项的 n 和百分比'),
            make_bullet('注：问卷为单次施测（非重复测量），不进行推断性统计（无组内变异可检验）'),
            # Output files LAST
            make_para('Normal', '输出文件（共 3 个 CSV + 2 个 PNG）：'),
            make_bullet('output/tables/13_post_survey_raw.csv  ← 参与者级别原始文本回答 + 数字编码'),
            make_bullet('output/tables/13_post_survey_descriptives.csv  ← 9 题的 M、SD、中位数、IQR、众数'),
            make_bullet('output/tables/13_post_survey_frequencies.csv  ← 每题各选项的 n 和百分比'),
            make_bullet('output/figures/fig_post_survey_stacked.png  ← 9 题堆叠水平条形图（含中位数虚线）'),
            make_bullet('output/figures/fig_post_survey_summary.png  ← 归一化点图：中位数（红圆）± IQR（蓝线）+ 均值（蓝菱形）'),
        ]

        for new_p in new_sections:
            ref_elem.addprevious(new_p)

        print(f"  Inserted §3.15 section ({len(new_sections)} paragraphs) before P{insert_before_idx:03d}")

# ─── Save ──────────────────────────────────────────────────────────────────────
doc.save(DOCX_PATH)
print("\nProcess.docx saved successfully.")
print("Summary of changes:")
print("  1. §5.1 TTC planned contrasts corrected (显著 → 不显著, p=.426)")
print("  2. CSV count updated: 59 → 62")
print("  3. PNG count updated: 17 → 21")
print("  4. 10_plots.R figure count updated: 14 → 16")
print("  5. Script count updated: 14 → 15 (14 R + 1 Python)")
print("  6. §3.15 added: 13_post_survey_descriptives.py documentation")
print("  7. Date updated: 2026-05-27 → 2026-05-28")
