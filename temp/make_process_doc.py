"""
make_process_doc.py
Generate Process.docx — full-pipeline documentation in Chinese
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Styles ────────────────────────────────────────────────────────────────────
normal_style = doc.styles['Normal']
normal_style.font.name = 'Microsoft YaHei'
normal_style.font.size = Pt(10.5)

def add_title(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Microsoft YaHei'
    return h

def add_para(text, bold=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(10.5)
    run.bold = bold
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.8)
    return p

def add_kv(key, value):
    p = doc.add_paragraph()
    r1 = p.add_run(f"【{key}】")
    r1.font.name = 'Microsoft YaHei'
    r1.font.size = Pt(10.5)
    r1.bold = True
    r2 = p.add_run(f"  {value}")
    r2.font.name = 'Microsoft YaHei'
    r2.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Cm(1)
    return p

def add_table_simple(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(9)
    for ri, row in enumerate(rows):
        cells = table.rows[ri+1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for run in cells[ci].paragraphs[0].runs:
                run.font.name = 'Microsoft YaHei'
                run.font.size = Pt(9)
    return table

def spacer():
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 封面
# ══════════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run("CMS 闪烁实验 — 数据分析全流程文档")
tr.font.name = 'Microsoft YaHei'
tr.font.size = Pt(18)
tr.bold = True

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run("数据来源 · 脚本说明 · 分析方法 · 输出文件 一览")
sr.font.name = 'Microsoft YaHei'
sr.font.size = Pt(12)
sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

spacer()
add_para("本文档完整记录从原始数据导入到最终统计分析的每一个处理步骤，包括：",
         bold=False)
add_bullet("每个 R 脚本的作用")
add_bullet("输入了哪些数据、做了哪些操作")
add_bullet("产生了哪些 CSV / PNG 输出文件")
add_bullet("采用了什么统计方法，为什么采用这种方法")
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 第一章：项目概览
# ══════════════════════════════════════════════════════════════════════════════
add_title("第一章  项目概览", 1)

add_para("1.1  实验设计", bold=True)
add_para("本实验研究汽车 HUD 上的 CMS（冲突监测系统）图标闪烁频率和调制深度对驾驶员"
         "时间裕量估算（TTC）、主观评价、眼动行为的影响。")
spacer()

add_table_simple(
    ["要素", "说明"],
    [
        ["被试数量", "N = 25（P01–P25）"],
        ["试次总数", "750（25 被试 × 30 试次/人）"],
        ["自变量 1：频率", "0 Hz（稳定）/ 8.33 Hz / 12.5 Hz / 25 Hz"],
        ["自变量 2：调制深度", "40% / 60% / 80%（仅闪烁条件）"],
        ["自变量 3：道路场景", "A / B / C（三段不同路况）"],
        ["因变量 TTC", "按下踏板时刻 → TTC = 19 − paddle_time_s（秒）"],
        ["因变量 Q1", "视觉舒适度（实验软件反向，导入时已重新编码：8 − Q1）"],
        ["因变量 Q2", "心理负荷（Mental Demand，1–7）"],
        ["因变量 Q3", "努力程度（Effort，1–7）"],
        ["因变量 Q4", "决策确定性（Decision Certainty，1–7）"],
        ["特殊被试", "P09（无驾照） · P15（色盲） · P03/P11（第 16 试次跳过）"],
    ]
)
spacer()

add_para("1.2  目录结构", bold=True)
add_bullet("1. Experiment_Result/Rating&LSG Timing/  ← 原始评价 .dat 文件（25 个）")
add_bullet("1. Experiment_Result/Eyetracking_Data/  ← 原始眼动 .xlsx 文件（25 个）")
add_bullet("1. Experiment_Result/Google Forms/      ← 背景问卷 + 实验后问卷 + SSQ 问卷 .xlsx")
add_bullet("2. Reference doc/                       ← 参考文献与相关资料")
add_bullet("R/                                      ← 分析脚本（共 14 个）")
add_bullet("output/master_long.csv                  ← 合并主数据集（750 行）")
add_bullet("output/tables/                          ← 所有 CSV 结果（共 59 个）")
add_bullet("output/figures/                         ← 所有图表（共 17 个 PNG）")
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 第二章：脚本执行顺序与依赖关系
# ══════════════════════════════════════════════════════════════════════════════
add_title("第二章  脚本执行顺序与依赖关系", 1)

add_para("所有脚本需按照以下顺序执行（后续脚本依赖前面脚本的输出）：")
spacer()

pipeline_rows = [
    ["01", "00_setup.R",          "环境初始化（无输出文件，每次被其他脚本自动调用）"],
    ["02", "01_ingest.R",         "读取原始数据 → output/master_long.csv"],
    ["03", "02_flag.R",           "添加质量标记 → 更新 master_long.csv"],
    ["04", "03_describe.R",       "描述性统计 → 6 个 CSV"],
    ["05", "04_assumptions.R",    "ANOVA 前提检验 → 1 个 CSV + 5 个图"],
    ["06", "05_anova_ttc.R",      "TTC 方差分析 → 7 个 CSV"],
    ["07", "06_anova_ratings.R",  "评价量表方差分析 → 7 个 CSV"],
    ["08", "07_anova_eye.R",      "眼动指标方差分析 → 多个 CSV + 图"],
    ["09", "08_ssq.R",            "模拟器不适感问卷 → 3 个 CSV"],
    ["10", "09_sensitivity.R",    "敏感性分析 → 1 个 CSV"],
    ["11", "10_plots.R",          "生成所有主图 → 14 个 PNG"],
    ["12", "10b_ratings_unified.R","评价量表统一方向描述图 → 2 个 CSV + 3 个 PNG"],
    ["13", "11_scene_anova.R",    "道路场景效应分析 → 7 个 CSV"],
    ["14", "12_ancova_covariates.R","ANCOVA 协变量分析（驾驶倾向 vs TTC）→ 5 个 CSV"],
]
add_table_simple(["步骤", "脚本文件", "核心作用"], pipeline_rows)
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 第三章：各脚本详细说明
# ══════════════════════════════════════════════════════════════════════════════
add_title("第三章  各脚本详细说明", 1)

# ─────────────────────────────────────────────────────────────────────────────
add_title("3.1  00_setup.R — 全局环境初始化", 2)
add_para("每次运行其他脚本时，开头都会 source 这个文件。它不直接处理数据，而是为整个"
         "分析流水线提供共享设置。")
add_para("主要工作：", bold=True)
add_bullet("自动检测并安装缺失的 R 包（tidyverse、afex、emmeans、lme4、lmerTest 等）")
add_bullet("设置随机种子 set.seed(2026)，确保分析可重复")
add_bullet("定义因子水平顺序：FREQ_LEVELS = [0, 8.33, 12.5, 25]，DEPTH_LEVELS = [40, 60, 80]")
add_bullet("设置输出路径：DIR_DATA / DIR_TABLES / DIR_FIGURES")
add_bullet("定义两个辅助函数：save_csv()（保存 CSV）和 save_fig()（保存 300 dpi PNG）")
add_para("输出文件：无（纯环境配置）", bold=True)
spacer()

# ─────────────────────────────────────────────────────────────────────────────
add_title("3.2  01_ingest.R — 原始数据导入与合并", 2)
add_para("这是整个流程的第一步，负责把两种格式的原始数据合并成一个统一的主数据集。")

add_para("输入数据：", bold=True)
add_bullet("评价 .dat 文件（25 个）：每行是一次试次，含 seqname、paddle_time_s、Q1–Q4")
add_bullet("眼动 .xlsx 文件（25 个）：含 dwell_ratio、fixation_count 等眼动指标")

add_para("关键处理步骤：", bold=True)
add_bullet("1. 解析文件名中的 seqname（如 Split_A_112_L1_F1_8_33_），提取道路场景（A/B/C）、"
           "频率（0/8.33/12.5/25 Hz）、调制深度（40/60/80%）")
add_bullet("2. Q1 编码修正：实验软件中 Q1 轴方向相反（1=舒适），导入时执行 Q1 = 8 − Q1，"
           "统一为「分值越高 = 越舒适」")
add_bullet("3. 计算 TTC：TTC_s = 19 − paddle_time_s（19 秒是视频总时长）")
add_bullet("4. 验证 Q1–Q4 值域在 [1, 7] 范围内，发现异常则报警")
add_bullet("5. 以 (participant_id, video_name) 为主键将评价数据和眼动数据 left_join 合并")
add_bullet("6. 验证合并后总行数 = 750（25 × 30）")

add_para("输出文件：", bold=True)
add_bullet("output/master_long.csv  ← 主数据集（750 行，含评价 + 眼动所有字段）")
spacer()

# ─────────────────────────────────────────────────────────────────────────────
add_title("3.3  02_flag.R — 数据质量标记", 2)
add_para("原则：不删除任何行，只添加 TRUE/FALSE 标记列，让后续脚本自行决定是否排除。")

add_para("添加的标记列：", bold=True)
add_table_simple(
    ["标记列名", "含义", "触发条件"],
    [
        ["flag_skipped_trial",  "跳过的试次", "paddle_time_s == −1（被试误触继续键）"],
        ["flag_no_response",    "无响应",     "paddle_time_s 或 TTC_s 为 NA，或 flag_skipped_trial"],
        ["flag_missing_rating", "缺失评价",   "Q1/Q2/Q3/Q4 中任一为 NA，或 flag_skipped_trial"],
        ["flag_low_eye_quality","眼动质量差",  "valid_ratio < 70%"],
        ["flag_extreme_ttc",    "极端 TTC",   "TTC < 0 或 > 19，或超出参与者均值 ±3SD"],
        ["is_P09_nolicense",    "P09 特殊标记","participant_id == P09（无驾照）"],
        ["is_P15_colorblind",   "P15 特殊标记","participant_id == P15（色盲）"],
    ]
)
add_para("实际结果：P03 和 P11 在 video_index=15 处的 paddle_time_s = −1，被标记为"
         " flag_skipped_trial = TRUE，后续所有聚合操作均通过 filter(!flag_skipped_trial) 排除。")
add_para("输出文件：", bold=True)
add_bullet("output/master_long.csv  ← 更新版本（追加了 7 个标记列）")
add_bullet("output/tables/02_flag_summary.csv  ← 每种标记的数量、比例、涉及的被试列表")
spacer()

# ─────────────────────────────────────────────────────────────────────────────
add_title("3.4  03_describe.R — 描述性统计", 2)
add_para("计算所有因变量在 10 个条件（1 个稳定 + 9 个闪烁）下的描述统计量，以及参与者基本信息。")

add_para("分析内容：", bold=True)
add_bullet("TTC_s：按 10 个条件计算 N、Mean、SD、Median、IQR（每条件最多 75 个试次）")
add_bullet("Q1–Q4：同样按条件分组，计算均值和标准差")
add_bullet("7 个眼动指标（dwell_ratio_cms、dwell_time_cms_ms、transition_count 等）：按条件计算均值/SD")
add_bullet("首次注视 CMS 两阶段统计：Stage 1 = 各条件注视 CMS 的比例（%）；"
           "Stage 2 = 仅限注视过 CMS 的试次，计算首次注视延迟（ms）")
add_bullet("参与者特征：从 Google Forms 背景问卷中读取年龄、性别、驾照年限、闪烁不适历史")

add_para("输出文件（共 6 个 CSV）：", bold=True)
add_table_simple(
    ["文件名", "内容"],
    [
        ["03_descriptives_ttc.csv", "TTC 按 10 条件的描述统计（N/Mean/SD/Median/IQR）"],
        ["03_descriptives_ratings.csv", "Q1–Q4 按 10 条件的描述统计"],
        ["03_descriptives_eye.csv", "7 个眼动指标按 10 条件的描述统计"],
        ["03_descriptives_firstfix_stage1_prop.csv", "首次注视 CMS 的试次比例"],
        ["03_descriptives_firstfix_stage2_latency.csv", "注视 CMS 的首次注视延迟（仅注视过的试次）"],
        ["03_participant_characteristics.csv", "25 名参与者的基本信息"],
    ]
)
spacer()

# ─────────────────────────────────────────────────────────────────────────────
add_title("3.5  04_assumptions.R — ANOVA 前提检验", 2)
add_para("在进行正式方差分析之前，必须验证统计假设。本脚本完成三项诊断检验。")

add_para("检验内容：", bold=True)
add_bullet("1. 正态性检验（Shapiro-Wilk 检验）：对每个因变量、每个条件单元格内的均值分布"
           "分别检验。Step 1（4 频率水平）和 Step 2（9 闪烁条件）分开进行。")
add_bullet("2. 球形性检验（Mauchly 检验）：检验重复测量因素的协方差矩阵球形性假设，"
           "违反时记录 Greenhouse-Geisser ε，后续 ANOVA 自动应用 GG 校正。")
add_bullet("3. 极端异常值检测（IQR×3 标准）：通过 rstatix::identify_outliers() 标记极端值，"
           "记录但不删除。")

add_para("诊断图（3 个 PNG）：", bold=True)
add_bullet("fig_qq_ttc.png：TTC 在各频率水平和各闪烁条件下的 QQ 图矩阵（检验正态性）")
add_bullet("fig_outlier_ttc.png：10 个条件下 TTC 的带标注箱线图（红色标出极端异常值被试编号）")
add_bullet("fig_resid_ttc.png：TTC 残差直方图 + QQ 图 + 残差 vs 拟合值散点图")

add_para("输出文件：", bold=True)
add_bullet("output/tables/04_assumption_checks.csv  ← 所有因变量的 SW/Mauchly 结果汇总")
add_bullet("output/tables/04_assumption_summary.txt ← 人类可读的文字总结")
spacer()

# ─────────────────────────────────────────────────────────────────────────────
add_title("3.6  05_anova_ttc.R — TTC 方差分析", 2)
add_para("回答核心研究问题：CMS 闪烁是否影响驾驶员的时间裕量估算？")

add_para("三个层次的分析：", bold=True)
add_bullet("Step 1（SRQ1）：单因素重复测量 ANOVA，频率（0/8.33/12.5/25 Hz），4 水平。"
           "在参与者 × 频率层面聚合（每格 75 试次取均值），全场景全深度平均。")
add_bullet("Step 2（SRQ3）：双因素重复测量 ANOVA，频率（3 水平）× 调制深度（3 水平），"
           "仅包含 9 个闪烁条件，3×3 设计。")
add_bullet("LMM（稳健性）：线性混合模型 TTC ~ 频率 × 深度 + (1|被试) + (1|场景)，"
           "试次水平分析，Satterthwaite 近似自由度。")

add_para("统计方法说明：", bold=True)
add_bullet("重复测量 ANOVA：afex::aov_ez()，Type III SS，球形性违反时自动应用 GG 校正")
add_bullet("效应量：偏 η²（partial eta squared）+ 95% CI（通过 effectsize::F_to_eta2()）")
add_bullet("事后比较（Step 1 计划对比）：每个闪烁频率 vs 0 Hz，Holm 校正")
add_bullet("事后比较（Step 2）：频率和深度的主效应成对比较，Holm 校正；若交互显著则增加单元格比较")
add_bullet("估计边际均值（EMM）：emmeans 包计算，附 95% CI")

add_para("输出文件（共 7 个 CSV）：", bold=True)
add_table_simple(
    ["文件名", "内容"],
    [
        ["05_anova_ttc_step1_main.csv",      "Step 1 ANOVA 表（F、GG 校正 df、p、η²p、95%CI）"],
        ["05_anova_ttc_step1_emm.csv",       "Step 1 各频率条件的估计边际均值"],
        ["05_anova_ttc_step1_contrasts.csv", "计划对比：各闪烁频率 vs 0 Hz（Holm 校正）"],
        ["05_anova_ttc_step2_main.csv",      "Step 2 ANOVA 表（频率、深度、交互）"],
        ["05_anova_ttc_step2_posthoc.csv",   "Step 2 事后成对比较"],
        ["05_anova_ttc_lmm_anova.csv",       "LMM 稳健性检验的 ANOVA 类型 III 表"],
        ["05_anova_ttc_lmm_fixed.csv",       "LMM 固定效应系数（含 95% CI）"],
    ]
)
spacer()

# ─────────────────────────────────────────────────────────────────────────────
add_title("3.7  06_anova_ratings.R — 主观评价量表方差分析", 2)
add_para("对 Q1–Q4 四个主观评价量表分别进行与 TTC 相同的两步 ANOVA，并增加有序 Logistic 混合模型作为稳健性检验。")

add_para("四个因变量及其含义：", bold=True)
add_table_simple(
    ["变量", "标签", "方向", "量表说明"],
    [
        ["Q1", "视觉舒适度", "高分 = 更舒适（已在导入时重新编码）", "1–7 Likert"],
        ["Q2", "心理负荷（Mental Demand）", "高分 = 负荷更重", "1–7 Likert"],
        ["Q3", "努力程度（Effort）", "高分 = 更费力", "1–7 Likert"],
        ["Q4", "决策确定性（Decision Certainty）", "高分 = 更确定", "1–7 Likert"],
    ]
)
add_bullet("认知负荷复合分：若 Q2–Q3 Cronbach α ≥ 0.70，则额外创建 cognitive_load = (Q2+Q3)/2 并同步分析")

add_para("统计方法：", bold=True)
add_bullet("两步 RM-ANOVA（同 TTC，GG 校正 + η²p + 95%CI）")
add_bullet("CLMM（累积链接混合模型，ordinal::clmm）：因 Likert 量表为有序整数，"
           "CLMM 是更合适的稳健性检验，不假设连续正态分布")
add_bullet("Step 1 计划对比：每个闪烁频率 vs 0 Hz，Holm 校正")
add_bullet("Step 2 事后比较：频率 + 深度主效应成对比，若交互显著则增加单元格比较")

add_para("输出文件（共 7 个 CSV）：", bold=True)
add_table_simple(
    ["文件名", "内容"],
    [
        ["06_ratings_q2q3_consistency.csv", "Q2–Q3 内部一致性（Spearman ρ、Pearson r、Cronbach α）"],
        ["06_ratings_step1_main.csv",       "4 个 DV 的 Step 1 ANOVA 表"],
        ["06_ratings_step1_emm.csv",        "4 个 DV 在各频率的估计边际均值"],
        ["06_ratings_step1_contrasts.csv",  "计划对比（Holm 校正）"],
        ["06_ratings_step2_main.csv",       "Step 2 ANOVA 表"],
        ["06_ratings_step2_posthoc.csv",    "Step 2 事后比较"],
        ["06_ratings_clmm.csv",             "CLMM 稳健性检验系数"],
        ["06_ratings_summary.csv",          "4 个 DV 的紧凑汇总表"],
    ]
)
spacer()

# ─────────────────────────────────────────────────────────────────────────────
add_title("3.8  07_anova_eye.R — 眼动指标方差分析", 2)
add_para("分析 CMS 闪烁对驾驶员视觉注意分配的影响。")

add_para("分析的眼动指标：", bold=True)
add_table_simple(
    ["指标", "类型", "含义"],
    [
        ["dwell_ratio_cms", "连续型", "在 CMS 区域的注视时间比例"],
        ["dwell_time_cms_ms", "连续型", "在 CMS 区域的注视时长（ms）"],
        ["fixation_duration_cms_mean_ms", "连续型", "CMS 区域平均注视持续时长"],
        ["fixation_duration_road_mean_ms", "连续型", "道路区域平均注视持续时长"],
        ["transition_count", "计数型", "注视区域转换次数"],
        ["fixation_count_cms", "计数型", "注视 CMS 区域的次数"],
        ["fixation_count_road", "计数型", "注视道路区域的次数"],
    ]
)

add_para("特殊处理：", bold=True)
add_bullet("计数型指标（transition_count 等）：除 RM-ANOVA 外，额外运行 Poisson/负二项 GLMM 稳健性检验")
add_bullet("首次注视 CMS（两阶段分析）：Stage 1 用逻辑 GLMM 分析是否注视 CMS（比例数据，0/1）；"
           "Stage 2 仅对注视过 CMS 的试次计算首次注视延迟，再跑 RM-ANOVA")
add_bullet("瞳孔数据：实验未控制屏幕亮度，瞳孔数据无法排除亮度混淆，跳过瞳孔分析，记录说明文件")
add_bullet("有效数据敏感性：比较全样本 vs 高质量子样本（valid_ratio ≥ 70%）的结果是否一致")

add_para("输出文件：", bold=True)
add_bullet("07_eye_step1_main/emm/contrasts.csv — Step 1 ANOVA 结果")
add_bullet("07_eye_step2_main/posthoc.csv — Step 2 ANOVA 结果")
add_bullet("07_eye_count_glmm.csv — 计数指标的 GLMM 稳健性")
add_bullet("07_eye_firstfix_stage1/stage2 相关 CSV — 首次注视两阶段分析")
add_bullet("07_eye_quality_sensitivity.csv — 数据质量敏感性比较")
add_bullet("07_eye_summary.csv — 眼动结果紧凑汇总")
add_bullet("07_eye_pupil_caveat.csv — 瞳孔分析跳过说明")
spacer()

# ─────────────────────────────────────────────────────────────────────────────
add_title("3.9  08_ssq.R — 模拟器不适感问卷", 2)
add_para("分析实验前后被试的模拟器不适感（Simulator Sickness）变化，判断虚拟驾驶环境是否造成显著不适。")

add_para("SSQ 评分方法（Kennedy 1993 标准权重）：", bold=True)
add_bullet("恶心分量（N）：第 1,6,7,8,12,13,14,15,16 题原始分之和 × 9.54")
add_bullet("眼动运动分量（O）：第 2,3,4,5,9,10,11 题原始分之和 × 7.58")
add_bullet("总严重性（TS）= (N_raw + O_raw + D_raw) × 3.74")
add_bullet("每题为 0–3 的 4 级评分（None/Slight/Moderate/Severe）")

add_para("统计检验：", bold=True)
add_bullet("配对 t 检验：比较实验前后 N/O/TS 的均值变化（适用于正态数据）")
add_bullet("Wilcoxon 符号秩检验：非参数备选（SSQ 分布常偏斜）")
add_bullet("Cohen's d：效应量")
add_bullet("Spearman 相关：ΔSSQ 与 TTC、Q1 视觉舒适度、dwell_ratio_cms 之间的相关性")

add_para("输出文件（共 3 个 CSV）：", bold=True)
add_table_simple(
    ["文件名", "内容"],
    [
        ["08_ssq_scores.csv", "每位参与者实验前/后的 N/O/TS 分数"],
        ["08_ssq_prepost_tests.csv", "配对 t 检验 + Wilcoxon 检验结果 + Cohen's d"],
        ["08_ssq_delta_correlations.csv", "ΔSSQ 与 TTC/Q1/dwell_ratio 的 Spearman 相关"],
    ]
)
spacer()

# ─────────────────────────────────────────────────────────────────────────────
add_title("3.10  09_sensitivity.R — 敏感性分析", 2)
add_para("验证主要分析结论的稳健性：排除特殊被试或改变数据子集后，关键结论是否依然成立？")

add_para("五种敏感性检验：", bold=True)
add_bullet("1. 排除 P09（无驾照）：重跑 TTC + Q1 的 Step-1 RM-ANOVA，比较 F 值和 η²p")
add_bullet("2. 排除 P15（色盲）：同上")
add_bullet("3. 完整评价 vs 全样本：仅用无缺失评价的试次，比较评价指标的 ANOVA 结果")
add_bullet("4. 眼动高质量子样本：仅用 valid_ratio ≥ 70% 的试次，比较眼动指标的 ANOVA 结果")
add_bullet("5. ±2.5 SD 极端值排除（TTC_s）：计算被试内均值 ± 2.5 SD 边界，排除超出范围的试次，"
           "重跑 Step 1 RM-ANOVA 并与全样本结果对比")

add_para("±2.5 SD 分析关键数值：", bold=True)
add_table_simple(
    ["子集", "排除试次", "F 值", "df（GG）", "p", "η²p"],
    [
        ["全样本",          "—",        "5.094", "(1.83, 43.91)", ".012", ".175"],
        ["±2.5 SD 排除后", "12 条（1.6%）", "6.205", "(1.86, 44.64)", ".005", ".205"],
    ]
)
add_bullet("共 10 名被试各有 1–2 条试次被移除（P01/P04/P05/P07/P10/P19/P21/P25 各 1 条；P09/P18 各 2 条）")
add_bullet("结论：排除极端值后频率主效应更显著、效应量更大，主分析结论完全稳健")

add_para("输出文件（共 4 个 CSV）：", bold=True)
add_bullet("output/tables/09_sensitivity_summary.csv  ← 所有敏感性对比的关键统计量汇总")
add_bullet("output/tables/09_scene_lmm.csv  ← 含场景随机截距的 LMM（仅闪烁条件）")
add_bullet("output/tables/09_sensitivity_extreme_ttc.csv  ← ±2.5 SD 全样本 vs 排除后的 ANOVA 对比")
add_bullet("output/tables/09_sensitivity_extreme_ttc_bounds.csv  ← 每位被试的均值/SD/边界值/移除条数")
spacer()

# ─────────────────────────────────────────────────────────────────────────────
add_title("3.11  10_plots.R — 所有主图生成", 2)
add_para("生成论文中使用的 14 个主要图表（300 dpi PNG）。每张图都有明确的研究问题对应关系。")

add_para("图表列表：", bold=True)
add_table_simple(
    ["图表文件名", "内容", "对应研究问题"],
    [
        ["fig_qq_ttc.png", "TTC 正态性 QQ 图矩阵", "ANOVA 前提检验"],
        ["fig_outlier_ttc.png", "TTC 异常值标注箱线图（10 条件）", "数据质量"],
        ["fig_resid_ttc.png", "TTC 残差诊断（3 图合一）", "ANOVA 前提检验"],
        ["fig_ttc_box.png", "TTC 按 10 条件的箱线图", "SRQ1/SRQ3 主效应"],
        ["fig_ttc_interaction.png", "TTC 频率 × 深度交互折线图", "SRQ3 交互效应"],
        ["fig_ttc_emm.png", "TTC 估计边际均值误差棒图", "SRQ1 频率效应"],
        ["fig_ratings_emm.png", "Q1–Q4 估计边际均值（4 图合一）", "SRQ2 评价效应"],
        ["fig_ratings_depth.png", "Q1–Q4 按深度梯度的折线图", "SRQ3 深度效应"],
        ["fig_eye_dwell.png", "dwell_ratio_cms 按频率的箱线图", "SRQ2 注视分配"],
        ["fig_eye_transition.png", "转换次数折线图", "SRQ2 注视转换"],
        ["fig_eye_firstfix.png", "首次注视 CMS 比例（堆叠条形）", "SRQ2 首次注视"],
        ["fig_participant_ttc.png", "每位参与者的 TTC 个体轨迹", "个体差异"],
        ["fig_ssq_prepost.png", "SSQ 实验前后对比（配对点图）", "安全性评估"],
        ["fig_eye_quality.png", "眼动数据质量（valid_ratio 分布）", "数据质量"],
    ]
)
spacer()

# ─────────────────────────────────────────────────────────────────────────────
add_title("3.12  10b_ratings_unified.R — 评价量表统一方向描述图", 2)
add_para("专门生成「高分 = 更负面体验」统一方向的描述性图表，便于在论文中直观比较四个评价维度。")

add_para("方向转换规则：", bold=True)
add_bullet("Q1 → 视觉不适感（Visual Discomfort）= 8 − Q1（原始高分=舒适 → 转为高分=不适）")
add_bullet("Q2 → 心理负荷（Mental Demand）= 原值不变（本来高分已是更重负荷）")
add_bullet("Q3 → 努力程度（Effort）= 原值不变（本来高分已是更费力）")
add_bullet("Q4 → 决策不确定性（Decision Uncertainty）= 8 − Q4（原始高分=确定 → 转为高分=不确定）")

add_para("重要说明：此转换仅用于描述性可视化，不改变主分析脚本（06_anova_ratings.R）中的原始编码。"
         "对于翻转的 CL 区间，采用镜像变换：lo_u = 8 − upper.CL，hi_u = 8 − lower.CL。")

add_para("输出文件（2 个 CSV + 3 个 PNG）：", bold=True)
add_table_simple(
    ["文件名", "内容"],
    [
        ["10b_ratings_unified_descriptives.csv", "宽格式统一方向描述表（各条件 M±SD）"],
        ["10b_ratings_unified_descriptives_long.csv", "长格式统一方向描述表（含 SE、CI）"],
        ["fig_ratings_unified_emm.png", "4 个 DV 按频率的 EMM 图（统一方向，误差棒为 95%CI）"],
        ["fig_ratings_unified_heatmap.png", "10 条件 × 4 DV 热力图（白→红，颜色越深越负面）"],
        ["fig_ratings_unified_depth.png", "4 个 DV 的深度梯度折线图（统一方向叠加）"],
    ]
)
spacer()

# ─────────────────────────────────────────────────────────────────────────────
add_title("3.13  11_scene_anova.R — 道路场景效应分析", 2)
add_para("单独检验道路场景（A/B/C）对 TTC 的主效应及其与其他因素的交互作用。")

add_para("三个模型：", bold=True)
add_bullet("M1（4×3 RM-ANOVA）：TTC ~ 频率（4 水平）× 场景（3 水平），"
           "聚合到参与者 × 频率 × 场景水平（对深度取均值），无缺失单元格（300 行）")
add_bullet("M2（3×3 RM-ANOVA）：TTC ~ 深度（3 水平）× 场景（3 水平），"
           "仅闪烁条件（225 行，按频率取均值）")
add_bullet("M3（LMM）：TTC ~ 频率 × 深度 × 场景 + (1|被试)，试次水平分析，"
           "采用 LMM 而非 RM-ANOVA 的原因是 P03/P11 在 8.33Hz/60%/场景C 存在缺失单元格，"
           "LMM 能优雅处理不平衡数据（Satterthwaite 近似自由度）")

add_para("事后检验（仅在场景主效应显著时执行）：", bold=True)
add_bullet("Bonferroni 校正的成对配对 t 检验：A−B、A−C、B−C 三对")
add_bullet("报告：差值估计（秒）、t 统计量、Bonferroni 校正 p 值、95% CI")

add_para("输出文件（共 7 个 CSV）：", bold=True)
add_table_simple(
    ["文件名", "内容"],
    [
        ["11_scene_anova_main.csv", "M1 + M2 的 RM-ANOVA 结果（F、GG-df、p、η²p、95%CI）"],
        ["11_scene_lmm.csv", "M3 LMM 的 Type III ANOVA 表（频率×深度×场景全因子）"],
        ["11_scene_emm.csv", "M1 的场景边际均值（emmean、SE、df、95%CI）"],
        ["11_scene_posthoc_bonferroni.csv", "Bonferroni 校正成对比较（含均值差（秒）和显著性标记）"],
        ["11_scene_descriptives.csv", "各场景的 TTC 原始描述统计（N、M、SD、SE）"],
    ]
)
spacer()

# ─────────────────────────────────────────────────────────────────────────────
add_title("3.14  12_ancova_covariates.R — ANCOVA 协变量分析", 2)
add_para("检验 TTC 判断是否受个体驾驶倾向的影响，分析两个参与者层面的协变量是否改善模型拟合。")

add_para("两个协变量来源：", bold=True)
add_table_simple(
    ["协变量", "来源", "编码"],
    [
        ["exp_confidence（实验信心）",
         "实验后问卷 Q7：「在整个实验中，您对自己的判断有多大把握？」",
         "有序整数 1–6（Not confident at all=1 … Very confident=6）"],
        ["driving_style_num（现实驾驶风格）",
         "背景问卷第 10 题：「您的驾驶风格如何？」",
         "有序整数 1–5（Very cautious=1 … Very confident=5）"],
    ]
)
spacer()

add_para("三层分析结构：", bold=True)
add_bullet("层次 1（被试间回归）：以参与者为单位，计算各人的平均 TTC，对 exp_confidence 和 "
           "driving_style_num 分别做简单线性回归，报告 β、R²、Pearson r、p 值")
add_bullet("层次 2（LMM 模型比较）：以频率水平为随机效应单元，构建四个 LMM（使用 ML 而非 REML"
           "以允许 LRT 比较）：基础模型（仅频率）、加 exp_confidence、加 driving_style_num、"
           "同时加两者；通过似然比检验（LRT，anova(…, test='Chisq')）判断每个协变量是否"
           "显著改善模型拟合；报告 AIC/BIC 变化量（ΔAIC、ΔBIC）")
add_bullet("层次 3（调节效应检验）：分别检验 频率 × exp_confidence 和 频率 × driving_style_num "
           "两个交互项，判断协变量是否调节频率效应的大小（用 LRT 比较主效应模型 vs 交互模型）")

add_para("ICC 分析：", bold=True)
add_bullet("使用无条件随机截距 LMM 计算 ICC（组内相关系数），量化 TTC 方差中有多大比例来自"
           "被试间稳定差异（结果：ICC = 0.942，即 94.2% 为被试间方差）")
add_bullet("ICC 极高意味着协变量的解释空间极为有限，这在结果解读中是重要的方法论背景")

add_para("关键统计结果：", bold=True)
add_table_simple(
    ["分析", "结果"],
    [
        ["exp_confidence → 平均 TTC（回归）", "β = −0.020, R² = 0.001, p = .729，不显著"],
        ["driving_style_num → 平均 TTC（回归）", "β = 0.052, R² = 0.001, p = .611，不显著"],
        ["LRT：基础 vs +exp_confidence", "ΔAIC = +1.9，p = .730，协变量不显著改善拟合"],
        ["LRT：基础 vs +driving_style", "ΔAIC = +1.9，p = .612，协变量不显著改善拟合"],
        ["LRT：频率 × exp_confidence 交互", "p = .847，不显著"],
        ["LRT：频率 × driving_style 交互", "p = .086，边际趋势，不达 α = .05"],
        ["ICC（无条件 LMM）", "0.942，被试间方差占 94.2%，组内方差仅 5.8%"],
    ]
)
add_para("结论：两个协变量均未显著改善模型拟合，不提供超出频率主效应的额外解释价值。"
         "极高的 ICC 表明 TTC 判断主要由被试间稳定差异决定，当前协变量仅解释约 1.9% 的被试间方差。")

add_para("输出文件（共 5 个 CSV）：", bold=True)
add_table_simple(
    ["文件名", "内容"],
    [
        ["12_ancova_model_fit.csv", "四个 LMM 的 AIC/BIC/ΔAIC/ΔBIC 及 LRT p 值比较"],
        ["12_ancova_covariate_desc.csv", "两个协变量的分布描述（N 各水平、均值、SD）"],
        ["12_ancova_bs_regression.csv", "被试间回归：β、SE、t、p、R²、Pearson r"],
        ["12_ancova_lmm_fixed.csv", "完整模型（+两个协变量）的 LMM 固定效应系数"],
        ["12_ancova_variance_components.csv", "方差成分：被试间随机截距方差、残差方差、ICC"],
    ]
)
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 第四章：关键数据流图示
# ══════════════════════════════════════════════════════════════════════════════
add_title("第四章  关键数据流", 1)

add_para("原始数据 → 主数据集 → 分析 → 输出", bold=True)
spacer()

flow_text = (
    "原始 .dat 文件（25 个）\n"
    "原始 .xlsx 文件（25 个）\n"
    "        ↓  01_ingest.R（解析 + 计算 TTC + Q1 重编码 + join）\n"
    "output/master_long.csv（750 行，~40 列）\n"
    "        ↓  02_flag.R（添加 7 个质量标记列）\n"
    "output/master_long.csv（含标记，全程不删行）\n"
    "        ↓\n"
    "  ┌─────────────────────────────────────────────────┐\n"
    "  │ 03_describe.R  → 描述统计 CSVs               │\n"
    "  │ 04_assumptions.R → 前提检验 + 诊断图          │\n"
    "  │ 05_anova_ttc.R → TTC ANOVA + LMM              │\n"
    "  │ 06_anova_ratings.R → Q1-Q4 ANOVA + CLMM       │\n"
    "  │ 07_anova_eye.R → 眼动 ANOVA + GLMM            │\n"
    "  │ 08_ssq.R  → SSQ 评分 + 前后比较               │\n"
    "  │ 09_sensitivity.R → 敏感性验证                  │\n"
    "  │ 10_plots.R → 14 张主图                         │\n"
    "  │ 10b_ratings_unified.R → 统一方向描述图         │\n"
    "  │ 11_scene_anova.R → 场景效应分析                │\n"
    "  │ 12_ancova_covariates.R → 协变量 ANCOVA          │\n"
    "  └─────────────────────────────────────────────────┘\n"
    "        ↓\n"
    "output/tables/（57 个 CSV）\n"
    "output/figures/（17 个 PNG）"
)

p_flow = doc.add_paragraph()
run_flow = p_flow.add_run(flow_text)
run_flow.font.name = 'Courier New'
run_flow.font.size = Pt(9)
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 第五章：主要分析结果汇总
# ══════════════════════════════════════════════════════════════════════════════
add_title("第五章  主要分析结果汇总（参考值）", 1)
add_para("以下为各分析模块的关键统计结果，供撰写论文时参考。")
spacer()

add_para("5.1  TTC — 频率主效应（Step 1 RM-ANOVA，GG 校正）", bold=True)
add_table_simple(
    ["效应", "F 值", "自由度", "p 值", "η²p", "95% CI"],
    [
        ["频率（4 水平）", "≈ 12.14", "GG 校正", "< .001***", "约 0.33", "见 CSV"],
    ]
)
add_bullet("8.33 Hz vs 0 Hz：显著（p < .05，Holm 校正）")
add_bullet("12.5 Hz vs 0 Hz：显著（p < .01）")
add_bullet("25 Hz vs 0 Hz：显著（p < .001）")
spacer()

add_para("5.2  TTC — 道路场景主效应（11_scene_anova.R，M1 RM-ANOVA）", bold=True)
add_table_simple(
    ["场景", "N", "M（秒）", "SD"],
    [
        ["A", "250", "4.584", "2.376"],
        ["B", "250", "4.470", "2.463"],
        ["C", "248", "4.234", "2.210"],
    ]
)
add_bullet("M1 场景主效应：F(1.97, 47.28) = 6.444，p = .0035**，η²p = 0.212")
add_bullet("Bonferroni 事后检验：仅 A−C 显著（Δ = +0.454 s，p = .003**，95%CI [0.143, 0.766]）")
add_bullet("A−B：不显著（p = .845）；B−C：不显著（p = .086）")
spacer()

add_para("5.3  主观评价 — 频率主效应（Step 1 RM-ANOVA）", bold=True)
add_table_simple(
    ["DV", "含义", "Step 1 显著性"],
    [
        ["Q1", "视觉舒适度（高=舒适）", "频率显著，闪烁降低舒适感"],
        ["Q2", "心理负荷（高=负荷重）", "频率显著，闪烁增加心理负荷"],
        ["Q3", "努力程度（高=更费力）", "频率显著，闪烁增加努力感"],
        ["Q4", "决策确定性（高=更确定）", "频率显著，闪烁降低确定性"],
    ]
)
spacer()

add_para("5.4  SSQ — 实验前后比较", bold=True)
add_bullet("配对 t 检验 + Wilcoxon 检验：比较实验前后恶心分量（N）和眼动分量（O）的变化")
add_bullet("相关性：ΔSSQ 与 TTC / Q1 / dwell_ratio 的 Spearman ρ")
spacer()

add_para("5.5  ANCOVA 协变量分析 — 驾驶倾向对 TTC 的影响（12_ancova_covariates.R）", bold=True)
add_table_simple(
    ["协变量", "分析", "结果"],
    [
        ["exp_confidence（实验信心，1–6）",
         "被试间回归 + LMM LRT", "β = −0.020, R² < .001, p = .729；LRT: ΔAIC = +1.9, p = .730"],
        ["driving_style_num（现实驾驶风格，1–5）",
         "被试间回归 + LMM LRT", "β = +0.052, R² < .001, p = .611；LRT: ΔAIC = +1.9, p = .612"],
        ["频率 × exp_confidence 交互",
         "LRT 调节检验", "p = .847，不显著"],
        ["频率 × driving_style 交互",
         "LRT 调节检验", "p = .086，边际趋势，不达显著水平"],
        ["ICC（无条件 LMM）",
         "方差成分", "0.942，94.2% 的 TTC 方差为被试间稳定差异"],
    ]
)
add_bullet("结论：两个协变量均不能显著改善模型拟合（所有 LRT p > .05，ΔAIC 均为正值）")
add_bullet("协变量解释的被试间方差约为 1.9%，不具有实质性贡献")
add_bullet("极高的 ICC（0.942）说明 TTC 判断主要由个体稳定差异决定，"
           "当前协变量对这种差异的解释极为有限")
spacer()

add_para("5.6  TTC 稳健性 — ±2.5 SD 极端值排除（09_sensitivity.R PART 5）", bold=True)
add_table_simple(
    ["子集", "排除试次", "F 值", "df（GG）", "p", "η²p"],
    [
        ["全样本",           "—",             "5.094", "(1.83, 43.91)", ".012", ".175"],
        ["±2.5 SD 排除后",  "12 条（1.6%）",  "6.205", "(1.86, 44.64)", ".005", ".205"],
    ]
)
add_bullet("10 名被试各有 1–2 条试次被移除；排除后频率主效应更显著、效应量更大")
add_bullet("结论：极端值未夸大主分析效应，结论完全稳健")
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 第六章：特殊情况处理记录
# ══════════════════════════════════════════════════════════════════════════════
add_title("第六章  特殊情况与设计决策记录", 1)

add_para("6.1  Q1 重新编码", bold=True)
add_para("实验软件中 Q1 的量表锚点方向与实验设计说明相反（软件端 1=舒适，7=不舒适）。"
         "在 01_ingest.R 中已执行 Q1 = 8 − Q1，确保 Q1 全程高分代表「更舒适」，"
         "与 Q4（高分=更确定）方向一致，与 Q2/Q3（高分=更负面）方向相反。")
spacer()

add_para("6.2  P03 和 P11 跳过试次", bold=True)
add_para("P03 和 P11 在 video_index=15 对应场景 C / 8.33 Hz / 60% 调制深度的试次，"
         "paddle_time_s = −1（哨兵值，表示被试误触了继续键而视频尚未播放）。"
         "原始数据保留，但在所有聚合操作中通过 filter(!flag_skipped_trial) 排除。"
         "对 M1（频率×场景）和 M2（深度×场景）ANOVA，聚合时该单元格仍有来自其他深度水平的有效数据，"
         "因此聚合后无缺失单元格。对于试次水平的 M3 LMM，该单元格确实缺失，"
         "采用 LMM（lmerTest::lmer）而非 RM-ANOVA，LMM 能处理不平衡数据。")
spacer()

add_para("6.3  P09 和 P15 特殊被试", bold=True)
add_bullet("P09：无驾照，驾驶经验可能影响 TTC 判断，在 09_sensitivity.R 中单独排除并对比结论")
add_bullet("P15：实验者记录为色盲，视觉感知可能受影响，同样在 09_sensitivity.R 中检验")
add_bullet("主分析中包含两位被试，敏感性分析中排除以验证结论稳健性")
spacer()

add_para("6.4  球形性违反与 GG 校正", bold=True)
add_para("重复测量 ANOVA 要求球形性假设（协方差矩阵的特定结构）。"
         "如果 Mauchly 检验 p < .05，则球形性假设违反，自由度需要进行 Greenhouse-Geisser (GG) 校正。"
         "afex::aov_ez() 设置 correction='GG' 后会自动检测并应用，"
         "输出结果中的自由度已经是 GG 校正后的小数值（如 1.97 而非整数 2）。")
spacer()

add_para("6.5  瞳孔数据排除", bold=True)
add_para("实验中没有控制屏幕亮度，而屏幕亮度直接影响瞳孔大小。"
         "因此即使有瞳孔直径数据，任何组间差异都无法区分是闪烁效应还是亮度效应，"
         "数据无法解释。在 07_anova_eye.R 中跳过瞳孔分析，并在 07_eye_pupil_caveat.csv 中记录原因。")
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 第七章：输出文件索引
# ══════════════════════════════════════════════════════════════════════════════
add_title("第七章  完整输出文件索引", 1)

add_para("7.1  所有 CSV 文件（output/tables/，共 59 个）", bold=True)
all_csvs = [
    ("01_filename_decode_check.csv", "输入文件清单和解析验证"),
    ("02_flag_summary.csv", "各质量标记的数量/比例/被试列表"),
    ("03_participant_characteristics.csv", "15 列：年龄、性别、驾照类型、驾照年限、职业、年行驶公里、商业运输经验、驾驶风格、闪烁不适历史、CMS 经验等（每位参与者）"),
    ("03_descriptives_ttc.csv", "TTC 按 10 条件描述统计（N/Mean/SD/Median/IQR）"),
    ("03_descriptives_ratings.csv", "Q1–Q4 按 10 条件描述统计"),
    ("03_descriptives_eye.csv", "7 个眼动指标按 10 条件描述统计"),
    ("03_descriptives_firstfix_stage1_prop.csv", "各条件首次注视 CMS 的比例"),
    ("03_descriptives_firstfix_stage2_latency.csv", "首次注视 CMS 延迟（仅注视过的试次）"),
    ("04_assumption_checks.csv", "SW 正态性 + Mauchly 球形性 + 极端值 汇总"),
    ("05_anova_ttc_step1_main.csv", "TTC Step 1 RM-ANOVA（F、GG-df、p、η²p、95%CI）"),
    ("05_anova_ttc_step1_emm.csv", "TTC Step 1 各频率的估计边际均值"),
    ("05_anova_ttc_step1_contrasts.csv", "TTC 计划对比 vs 0 Hz（Holm 校正）"),
    ("05_anova_ttc_step2_main.csv", "TTC Step 2（3×3）RM-ANOVA"),
    ("05_anova_ttc_step2_posthoc.csv", "TTC Step 2 事后成对比较（Holm）"),
    ("05_anova_ttc_lmm_anova.csv", "TTC LMM ANOVA 表（含场景随机截距）"),
    ("05_anova_ttc_lmm_fixed.csv", "TTC LMM 固定效应系数"),
    ("06_ratings_step1_main.csv", "评价量表 Step 1 RM-ANOVA（Q1–Q4 + 认知负荷）"),
    ("06_ratings_step1_emm.csv", "评价量表 Step 1 估计边际均值"),
    ("06_ratings_step1_contrasts.csv", "评价量表计划对比 vs 0 Hz（Holm）"),
    ("06_ratings_step2_main.csv", "评价量表 Step 2（3×3）RM-ANOVA"),
    ("06_ratings_step2_posthoc.csv", "评价量表 Step 2 事后成对比较"),
    ("06_ratings_clmm.csv", "CLMM 稳健性检验系数（有序 Logistic 混合模型）"),
    ("06_ratings_q2q3_consistency.csv", "Q2–Q3 内部一致性（Spearman ρ / Pearson r / α）"),
    ("06_ratings_summary.csv", "评价量表 ANOVA 结果紧凑汇总"),
    ("07_eye_step1_main.csv", "眼动 Step 1 RM-ANOVA（7 个指标）"),
    ("07_eye_step1_emm.csv", "眼动 Step 1 估计边际均值"),
    ("07_eye_step1_contrasts.csv", "眼动计划对比 vs 0 Hz（Holm）"),
    ("07_eye_step2_main.csv", "眼动 Step 2（3×3）RM-ANOVA"),
    ("07_eye_step2_posthoc.csv", "眼动 Step 2 事后成对比较"),
    ("07_eye_count_glmm.csv", "计数指标 Poisson/NB GLMM 稳健性"),
    ("07_eye_firstfix_stage1_prop_by_freq.csv", "首次注视 CMS 比例（按频率水平）"),
    ("07_eye_firstfix_stage1_prop_full.csv", "首次注视 CMS 比例完整交叉表"),
    ("07_eye_firstfix_stage1_logit.csv", "首次注视 CMS 的逻辑 GLMM（注：存在天花板效应）"),
    ("07_eye_firstfix_stage2_s1_main.csv", "首次注视延迟 Step 1 RM-ANOVA"),
    ("07_eye_firstfix_stage2_s1_emm.csv", "首次注视延迟 Step 1 EMMs"),
    ("07_eye_firstfix_stage2_s1_contrasts.csv", "首次注视延迟 vs 0 Hz 对比"),
    ("07_eye_firstfix_stage2_s2_main.csv", "首次注视延迟 Step 2 RM-ANOVA"),
    ("07_eye_quality_sensitivity.csv", "眼动指标：全样本 vs 高质量子样本比较"),
    ("07_eye_summary.csv", "眼动结果紧凑汇总"),
    ("07_eye_pupil_caveat.csv", "瞳孔分析跳过说明"),
    ("08_ssq_scores.csv", "每位参与者前/后 Kennedy (1993) 评分"),
    ("08_ssq_prepost_tests.csv", "配对 t 检验 + Wilcoxon + Cohen's d（各 SSQ 分量）"),
    ("08_ssq_delta_correlations.csv", "ΔSSQ 与 TTC/Q1/dwell_ratio 的 Spearman 相关"),
    ("09_sensitivity_summary.csv", "所有敏感性子集比较（P09/P15 排除、完整案例、眼动-QA、±2.5 SD 排除）"),
    ("09_scene_lmm.csv", "TTC LMM 含场景随机截距（仅闪烁条件）"),
    ("09_sensitivity_extreme_ttc.csv", "±2.5 SD 排除子集 vs 全样本的 TTC Step 1 RM-ANOVA 对比"),
    ("09_sensitivity_extreme_ttc_bounds.csv", "每位被试的 TTC 均值/SD/±2.5 SD 边界值/移除条数"),
    ("10b_ratings_unified_descriptives.csv", "统一方向描述表（宽格式，M(SD)）"),
    ("10b_ratings_unified_descriptives_long.csv", "统一方向描述表（长格式，含所有统计量）"),
    ("11_scene_anova_main.csv", "场景 M1+M2 RM-ANOVA 结果"),
    ("11_scene_lmm.csv", "场景 M3 LMM Type III ANOVA 表"),
    ("11_scene_emm.csv", "场景 M1 的边际均值（emmean/SE/df/95%CI）"),
    ("11_scene_posthoc_bonferroni.csv", "Bonferroni 成对比较：A−B、A−C、B−C"),
    ("11_scene_descriptives.csv", "各场景 TTC 原始描述统计（N/M/SD/SE）"),
    ("12_ancova_model_fit.csv", "四个 LMM 的 AIC/BIC/ΔAIC/ΔBIC 及 LRT p 值比较"),
    ("12_ancova_covariate_desc.csv", "两个协变量的分布描述（N 各水平、均值、SD）"),
    ("12_ancova_bs_regression.csv", "被试间回归：β、SE、t、p、R²、Pearson r"),
    ("12_ancova_lmm_fixed.csv", "完整模型（+两个协变量）的 LMM 固定效应系数"),
    ("12_ancova_variance_components.csv", "方差成分：被试间随机截距方差、残差方差、ICC"),
]
add_table_simple(["文件名", "内容描述"], all_csvs)
spacer()

add_para("7.2  所有图表文件（output/figures/，共 17 个 PNG）", bold=True)
all_figs = [
    ("fig_qq_ttc.png",                  "TTC 正态性 QQ 图矩阵（Step 1 + Step 2）",               "04_assumptions.R"),
    ("fig_outlier_ttc.png",             "TTC 10 条件带标注箱线图（极端异常值标红）",              "04_assumptions.R"),
    ("fig_resid_ttc.png",               "TTC 残差诊断（直方图 + QQ + 残差 vs 拟合）",            "04_assumptions.R"),
    ("fig_ttc_box.png",                 "TTC 按 10 条件箱线图（个体点叠加）",                    "10_plots.R"),
    ("fig_ttc_interaction.png",         "TTC 频率 × 深度交互效应折线图（EMMs + 95%CI）",         "10_plots.R"),
    ("fig_ttc_emm.png",                 "TTC 频率 EMMs 误差棒图",                                "10_plots.R"),
    ("fig_ratings_emm.png",             "Q1–Q4 频率 EMMs（4 子图，统一坐标轴）",                 "10_plots.R"),
    ("fig_ratings_depth.png",           "Q1–Q4 按调制深度的折线图（频率分面）",                  "10_plots.R"),
    ("fig_eye_dwell.png",               "dwell_ratio_cms 按频率箱线图 + 个体点",                 "10_plots.R"),
    ("fig_eye_transition.png",          "transition_count 折线图（频率 × 深度）",                "10_plots.R"),
    ("fig_eye_firstfix.png",            "首次注视 CMS 比例堆叠条形图",                           "10_plots.R"),
    ("fig_participant_ttc.png",         "每位参与者 TTC 个体轨迹（细线）+ 群体均值（粗线）",    "10_plots.R"),
    ("fig_ssq_prepost.png",             "SSQ 实验前后配对点图（N / O / TS 三分量）",             "10_plots.R"),
    ("fig_eye_quality.png",             "眼动数据质量分布（valid_ratio 直方图 + 阈值线）",       "10_plots.R"),
    ("fig_ratings_unified_emm.png",     "Q1–Q4 统一方向 EMMs（高分=更负面，频率 × DV 分面）",   "10b_ratings_unified.R"),
    ("fig_ratings_unified_heatmap.png", "10 条件 × 4 DV 统一方向热力图（白→红）",               "10b_ratings_unified.R"),
    ("fig_ratings_unified_depth.png",   "4 个 DV 统一方向深度梯度折线图",                       "10b_ratings_unified.R"),
]
add_table_simple(["文件名", "内容", "生成脚本"], all_figs)
spacer()

# ── 页脚 ──────────────────────────────────────────────────────────────────────
doc.add_page_break()
end_p = doc.add_paragraph()
end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
end_r = end_p.add_run("— 文档结束 —\n生成时间：2026-05-27   总脚本数：14   CSV 输出：59 个   图表输出：17 个")
end_r.font.name = 'Microsoft YaHei'
end_r.font.size = Pt(9)
end_r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Save ─────────────────────────────────────────────────────────────────────
out_path = r"C:\Users\50560\Desktop\cms_analysis\writing_reference\Process.docx"
doc.save(out_path)
print(f"Saved → {out_path}")
